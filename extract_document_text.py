"""Document text extraction with optional Docling integration.

When Docling is installed (``pip install docling``), rich extraction is used
for PDFs (tables, figures, OCR), DOCX, PPTX, HTML, and images.  When Docling
is not available, the script falls back to basic extraction using ``pypdf``
and ``python-docx``.

CLI interface (unchanged from previous versions):
    python extract_document_text.py <file_path>

Output: a single JSON line to stdout with the shape:
    { "ok": true, "text": "...", "metadata": { ... }, "tables": [...], "figures": [...] }
    or
    { "ok": false, "error": "..." }
"""

import json
import sys
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Docling-based extraction (rich)
# ---------------------------------------------------------------------------

_HAS_DOCLING = False
try:
    from docling.document_converter import DocumentConverter  # type: ignore[import-untyped]

    _HAS_DOCLING = True
except ImportError:
    pass


def _extract_with_docling(path: Path) -> dict[str, Any]:
    """Use Docling for rich extraction supporting PDF, DOCX, PPTX, HTML, images."""
    converter = DocumentConverter()
    result = converter.convert(str(path))
    doc = result.document

    text = doc.export_to_markdown()

    tables: list[dict[str, Any]] = []
    for table in doc.tables:
        table_data = table.export_to_dataframe()
        tables.append(
            {
                "headers": list(table_data.columns),
                "rows": table_data.values.tolist(),
            }
        )

    figures: list[dict[str, str]] = []
    for picture in doc.pictures:
        parts: list[str] = []
        for cap in picture.captions:
            if cap.text:
                parts.append(cap.text)
        description = " ".join(parts) if parts else "Figure (no caption)"
        figures.append({"description": description})

    metadata: dict[str, Any] = {
        "extractor": "docling",
        "format": path.suffix.lower().lstrip("."),
        "table_count": len(tables),
        "figure_count": len(figures),
    }

    return {
        "ok": True,
        "text": text[:50000],
        "metadata": metadata,
        "tables": tables,
        "figures": figures,
    }


# ---------------------------------------------------------------------------
# Fallback extraction (basic)
# ---------------------------------------------------------------------------


def _extract_pdf_basic(path: Path) -> str:
    from pypdf import PdfReader  # type: ignore[import-untyped]

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages[:25]:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx_basic(path: Path) -> str:
    import docx  # type: ignore[import-untyped]

    document = docx.Document(str(path))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(parts)


def _extract_with_fallback(path: Path) -> dict[str, Any]:
    """Basic extraction using pypdf / python-docx for supported formats."""
    extension = path.suffix.lower()

    if extension in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif extension == ".pdf":
        text = _extract_pdf_basic(path)
    elif extension == ".docx":
        text = _extract_docx_basic(path)
    else:
        raise ValueError(f"Unsupported extension: {extension}")

    metadata: dict[str, Any] = {
        "extractor": "fallback",
        "format": extension.lstrip("."),
        "table_count": 0,
        "figure_count": 0,
    }

    return {
        "ok": True,
        "text": text[:50000],
        "metadata": metadata,
        "tables": [],
        "figures": [],
    }


# ---------------------------------------------------------------------------
# Unified extraction entry point
# ---------------------------------------------------------------------------

# Extensions that Docling supports beyond our basic fallback set.
_DOCLING_EXTENSIONS = {".pdf", ".docx", ".pptx", ".html", ".htm", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
_BASIC_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def extract(path: Path) -> dict[str, Any]:
    """Extract text (and optionally tables/figures) from a document file.

    Uses Docling when available and the format is supported; otherwise falls
    back to the basic pypdf / python-docx extractors.
    """
    extension = path.suffix.lower()

    if _HAS_DOCLING and extension in _DOCLING_EXTENSIONS:
        return _extract_with_docling(path)

    if extension in _BASIC_EXTENSIONS:
        return _extract_with_fallback(path)

    raise ValueError(f"Unsupported extension: {extension}")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------


def main() -> None:
    if len(sys.argv) != 2:
        print(json.dumps({"ok": False, "error": "Expected a single file path argument."}))
        return

    path = Path(sys.argv[1])
    try:
        result = extract(path)
        print(json.dumps(result, ensure_ascii=False, default=str))
    except Exception as exc:  # pragma: no cover
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
